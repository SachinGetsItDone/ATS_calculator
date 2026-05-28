# Standard library: in-memory byte streams (used to treat raw bytes like a file)
import io
# Third-party: detects the true file type from raw bytes (more reliable than file extension)
import magic
# Type hints: Tuple for multi-value returns, Optional for nullable values
from typing import Tuple, Optional, Tuple

# pdfplumber: primary PDF text extractor — handles most modern PDFs well
import pdfplumber
# python-docx: extracts text from .docx Word documents
from docx import Document
# PyPDF2: fallback PDF extractor — used when pdfplumber fails
import PyPDF2

# Custom error classes and logging/utility helpers from our utils module
from backend.utils.file_utils import(
    FileParsingError,       # raised when a file can't be parsed
    TextExtractionError,    # raised when text extraction yields nothing
    FileUploadError,        # raised for upload-level failures
    log_error,              # logs exceptions with context label
    log_warning,            # logs non-fatal warnings with context label
    log_info,               # logs informational messages with context label
    with_fallback           # runs fn A, falls back to fn B on failure
)

# Config constants: size limits and allowed MIME types (e.g. application/pdf)
from backend.core.config import (
    MAX_FILE_SIZE_BYTES,    # max upload size in bytes
    MAX_FILE_SIZE_MB,       # same limit in MB (used in user-facing messages)
    SUPPORTED_MIME_TYPES    # dict mapping MIME type → file_type string (e.g. 'pdf')
)

# Raised when the file itself is corrupt or unparseable
class FileParsingError(Exception):
    pass

# Raised when the file fails validation (wrong type, too big, empty, etc.)
class FileValidationError(Exception):
    pass

# Checks that the uploaded file is safe to process before we touch its contents
# Returns: (is_valid, error_message, file_type_string)
def validate_file(file_data:bytes, filename:str)->Tuple[bool, str, Optional[str]]:
    # Count the raw bytes to determine file size
    file_size_bytes = len(file_data)
    # Reject if file is larger than our configured maximum
    if file_size_bytes > MAX_FILE_SIZE_BYTES:
        # Convert bytes → MB for a readable error message
        size_mb = file_size_bytes / (1024 * 1024)
        return False, (
            f'File size ({size_mb:.2f} MB) exceeds the maximum of {MAX_FILE_SIZE_MB} MB. '
            'Please upload a smaller file or compress your resume.'
        ), None

    # Reject empty files — nothing to extract from them
    if file_size_bytes==0:
        return False, 'uploade file is empty...please check the file you have uploaded and try again'

    try:
        # Use libmagic to sniff the real MIME type from the file's byte signature
        # (safer than trusting the filename extension, which can be spoofed)
        mime_type=magic.from_buffer(file_data, mime=True)
    except Exception as e:
        # If magic itself crashes, we can't safely identify the file
        return False, f"error deteminin the file type : {e}", None

    # Reject MIME types we don't support (e.g. images, zip files, etc.)
    if mime_type not in SUPPORTED_MIME_TYPES:
        # Build a human-readable list of allowed formats for the error message
        supported=', '.join(SUPPORTED_MIME_TYPES.keys()).upper()
        return False, (
            f'Unsupported file type: {mime_type}. '
            f'Please upload one of: {supported}.'
        ), None

    # All checks passed — return True with no error and the resolved file type string
    return True, '', SUPPORTED_MIME_TYPES[mime_type]

# Pulls clickable hyperlinks embedded in a PDF (e.g. LinkedIn, GitHub URLs)
# These live in annotation objects, not the visible text layer
def _extract_pdf_hyperlinks(file_data: bytes) -> str:
    # Collect all discovered URLs here
    urls = []
    try:
        # Wrap raw bytes in a file-like object so PyPDF2 can read it
        reader = PyPDF2.PdfReader(io.BytesIO(file_data))
        # Walk every page looking for annotation dictionaries
        for page in reader.pages:
            # Skip pages that have no annotations at all
            if '/Annots' not in page:
                continue
            # Each annotation is stored as an indirect reference — resolve it
            for annot_ref in page['/Annots']:
                try:
                    annot = annot_ref.get_object()
                    # We only care about Link annotations, not highlights/comments
                    if annot.get('/Subtype') != '/Link':
                        continue
                    # The action dict holds the destination URI for link annotations
                    action = annot.get('/A', {})
                    uri = action.get('/URI', '')
                    # Only keep non-empty string/bytes values
                    if uri and isinstance(uri, (str, bytes)):
                        # PyPDF2 may return bytes for URI values
                        if isinstance(uri, bytes):
                            # Decode bytes → string, drop any undecodable characters
                            uri = uri.decode('utf-8', errors='ignore')
                        # Strip surrounding whitespace
                        uri = uri.strip()
                        # Only keep full web URLs (ignore mailto:, tel:, etc.)
                        if uri.startswith('http'):
                            urls.append(uri)
                except Exception:
                    # Skip any malformed annotation silently
                    pass
    except Exception:
        # If the whole PDF is unreadable by PyPDF2, just return empty
        pass
    # Return all found URLs joined by newlines (or empty string if none)
    return '\n'.join(urls)


# Primary PDF text extractor using pdfplumber (handles most modern PDFs)
def _extract_pdf_with_pdfplumber(file_data: bytes) -> str:
    text = ''
    # Open the PDF from an in-memory byte stream
    with pdfplumber.open(io.BytesIO(file_data)) as pdf:
        # Iterate over every page and accumulate its text
        for page in pdf.pages:
            page_text = page.extract_text()
            # Only append if the page actually had readable text
            if page_text:
                text += page_text + '\n'

    # If we got nothing at all, raise so with_fallback can try PyPDF2 instead
    if not text.strip():
        raise TextExtractionError(
            'pdfplumber extracted no text',
            user_message='No text could be extracted from the PDF.'
        )

    # Also grab any embedded hyperlinks and append them after the body text
    hyperlinks = _extract_pdf_hyperlinks(file_data)
    if hyperlinks:
        text = text.strip() + '\n' + hyperlinks

    # Return clean text with no leading/trailing whitespace
    return text.strip()


# Fallback PDF text extractor using PyPDF2 (used when pdfplumber fails)
def _extract_pdf_with_pypdf2(file_data: bytes) -> str:
    text = ''
    # Open the PDF from bytes; PyPDF2 needs a file-like object
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
    # Walk every page and collect text
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + '\n'

    # Nothing extracted → raise so the caller knows this fallback also failed
    if not text.strip():
        raise TextExtractionError(
            'PyPDF2 extracted no text',
            user_message='No text could be extracted from the PDF.'
        )

    # Append any hyperlinks found in the PDF annotations
    hyperlinks = _extract_pdf_hyperlinks(file_data)
    if hyperlinks:
        text = text.strip() + '\n' + hyperlinks

    return text.strip()


# Public entry point for PDF extraction — tries pdfplumber first, PyPDF2 second
def extract_text_from_pdf(file_data: bytes) -> str:
    try:
        # with_fallback runs _extract_pdf_with_pdfplumber(file_data);
        # if that raises, it retries with _extract_pdf_with_pypdf2(file_data)
        # and returns (result, used_fallback=True)
        result, used_fallback=with_fallback(
        _extract_pdf_with_pdfplumber,
        _extract_pdf_with_pypdf2,
        file_data,
        log_fallback=True     # emit a warning log when the fallback is triggered
    )

        # Let the caller know we had to use the secondary extractor
        if used_fallback:
            log_info('PDF EXTRACTION succeded using the PyPDF2 fallback', context='resume_parser')
        return result

    except Exception as e:
        # Both extractors failed — log and surface a user-friendly error
        log_error(e, context='extract_text_from_pdf')
        raise FileParsingError(
            'Failed to extract text from PDF using both pdfplumber and PyPDF2. '
            'The PDF may be corrupted, password-protected, or contain only scanned images. '
            'Please ensure it contains selectable text.'
        ) from e


# Extracts text from a .docx Word document (paragraphs + table cells + hyperlinks)
def extract_text_from_docx(file_data: bytes) -> str:
    try:
        # Parse the DOCX from an in-memory byte stream
        doc = Document(io.BytesIO(file_data))
        # Accumulate all non-empty text chunks here
        text_parts = []

        # Walk top-level paragraphs (headings, bullets, body text, etc.)
        for paragraph in doc.paragraphs:
            # Skip blank paragraphs — they're just whitespace
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Walk every table → every row → every cell to capture tabular content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Skip empty cells
                    if cell.text.strip():
                        text_parts.append(cell.text)

        # Merge all pieces into one block of text separated by newlines
        text = '\n'.join(text_parts)

        # If nothing was extracted, the doc is empty or unreadable
        if not text.strip():
            raise FileParsingError(
                'No text could be extracted from the document. '
                'The document may be empty or corrupted.'
            )

        try:
            # Iterate the document's relationship table to find hyperlinks
            # (DOCX stores links as relationships, not inline text)
            for rel in doc.part.rels.values():
                if 'hyperlink' in rel.reltype.lower():
                    url = rel._target   # the actual URL string
                    # Only append absolute web URLs
                    if isinstance(url, str) and url.startswith('http'):
                        text += '\n' + url
        except Exception:
            # Non-fatal — skip hyperlinks if the rels table is unreadable
            pass

        log_info(f'Extracted {len(text)} chars from DOCX', context='resume_parser')
        # Return text with no leading/trailing whitespace
        return text.strip()

    except FileParsingError:
        raise   # Re-raise unchanged — don't wrap in another FileParsingError

    except Exception as e:
        log_error(e, context='extract_text_from_docx')
        raise FileParsingError(
            'Failed to extract text from DOCX. '
            'The document may be corrupted or in an unsupported format. '
            'Please try re-saving or converting to PDF.'
        ) from e

# Legacy .doc format handler — we don't support it, so always raise with guidance
def extract_text_from_doc(file_data: bytes) -> str:
    raise FileParsingError(
        'Legacy .doc format is not supported. '
        'Please convert your document to .docx or .pdf and try again. '
        'You can convert using Microsoft Word, Google Docs, or online tools.'
    )

# Routes extraction to the right function based on the detected file type
def extract_text(file_data:bytes, file_type:str)->str:
    if file_type=='pdf':
        # PDF path — tries pdfplumber then PyPDF2
        return extract_text_from_pdf(file_data)
    elif file_type=='docx':
        # Modern Word document path
        return extract_text_from_docx(file_data)
    elif file_type=='doc':
        # Old Word format — unsupported, raises immediately
        return extract_text_from_doc(file_data)
    else:
        # Anything else that slipped past validate_file
        raise FileValidationError(
            f'invalid file type: {file_type}. supported types are: pdf, docx and doc'


        )

# Top-level function called by the API layer — validates then extracts, returns text + metadata
def parse_resume_file(file_data: bytes, filename:str)->Tuple[str, dict]:
    log_info(f'parsing file :{filename}', context='parse_Resume_file')

    # --- Phase 1: Validate the file before touching its contents ---
    try:
        # validate_file returns (is_valid, error_msg, file_type)
        is_valid, error_msg, file_type=validate_file(file_data, filename)
        # If validation failed, surface the human-readable error to the caller
        if not is_valid:
            log_warning(f'valiudation failed for file {filename}', context='parse_resume_file')
            raise FileValidationError(error_msg)

    except FileValidationError as e:
        raise   # Bubble up unchanged — already a clean user-facing error

    except Exception as e:
        # Unexpected crash during validation (e.g. libmagic missing)
        log_error(e, context='parse_resume_file_validation')
        raise FileValidationError(
            'Could not validate the uploaded file. Please ensure it is a valid PDF or DOCX.'
        ) from e

    # --- Phase 2: Extract the text from the validated file ---

    try:
        # Dispatch to the correct extractor based on file_type ('pdf', 'docx', 'doc')
        text = extract_text(file_data, file_type)
        log_info(f'Extracted {len(text)} chars from {filename}', context='parse_resume_file')

    except FileParsingError:
        raise   # Re-raise unchanged

    except Exception as e:
        # Unexpected crash during extraction
        log_error(e, context='parse_resume_file_extraction')
        raise FileParsingError(
            'An unexpected error occurred while processing the file. '
            'Please try again or contact support if the problem persists.'
        ) from e

    # Build a summary dict with key facts about what was processed
    metadata = {
        'filename':        filename,        # original uploaded filename
        'file_type':       file_type,       # resolved type: 'pdf' or 'docx'
        'file_size_bytes': len(file_data),  # raw size of the upload
        'text_length':     len(text),       # number of characters extracted
        'success':         True,            # flag indicating clean extraction
    }
    # Return the extracted text alongside the metadata dict
    return text, metadata
